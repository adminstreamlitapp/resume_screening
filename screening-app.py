import os
import docx
from PyPDF2 import PdfReader
import streamlit as st

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + "\n"
    
    return text

def extract_text_from_pdf(file_path):
    text = ""
    with open(file_path, "rb") as f:
        reader = PdfReader(f)
        for page in reader.pages:
            text += page.extract_text()
    return text

def extract_text_from_file(file_path):
    if file_path.endswith('.docx'):
        return extract_text_from_docx(file_path)
    elif file_path.endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    else:
        return ""

def evaluate_resume(resume_text):
    required_skills = {
        "Core Java": False,
        "Java 8": False,
        "Spring": False,
        "Spring Boot": False,
        "Spring Data": False,
        "JPA": False
    }
    
    optional_skills = {
        "JMS Kafka": False,
        "JWT": False,
        "JWE Encryption": False
    }
    
    for skill in required_skills:
        if skill in resume_text:
            required_skills[skill] = True
            
    for skill in optional_skills:
        if skill in resume_text:
            optional_skills[skill] = True
    
    required_match_percentage = sum(required_skills.values()) / len(required_skills) * 100
    optional_match_percentage = sum(optional_skills.values()) / len(optional_skills) * 100
    
    return required_skills, optional_skills, required_match_percentage, optional_match_percentage

st.title('Resume Analysis')

# Input field for resumes directory
resumes_directory = st.text_input("Enter directory containing resumes:")

if st.button("Analyze Resumes"):
    if resumes_directory:
        for resume_file_name in os.listdir(resumes_directory):
            if not resume_file_name.startswith('~$'):
                resume_path = os.path.join(resumes_directory, resume_file_name)
                if os.path.isfile(resume_path) and (resume_file_name.endswith('.docx') or resume_file_name.endswith('.pdf')):
                    resume_text = extract_text_from_file(resume_path)
                    required_skills, optional_skills, required_match_percentage, optional_match_percentage = evaluate_resume(resume_text)

                    st.subheader(f"Results for resume: {resume_file_name}")

                    st.subheader("Required skills:")
                    for skill, present in required_skills.items():
                        st.write(f"{skill} : {'Yes' if present else 'No'}")

                    st.subheader("Optional Skills:")
                    for skill, present in optional_skills.items():
                        st.write(f"{skill} : {'Yes' if present else 'No'}")

                    st.write("Match % for Required skills:", required_match_percentage)
                    st.write("Match % for Optional Skills:", optional_match_percentage)
                    st.write("-" * 50)
    else:
        st.warning("Please enter resumes directory")
